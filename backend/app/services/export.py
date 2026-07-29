from app.models.domain_models import ConsolidatedResult, ModelResult
from app.utils.mdformatter import clean_markdown
from io import BytesIO
from datetime import datetime
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import ParagraphStyle

class Export:

    def __init__(self):
        pdfmetrics.registerFont(TTFont("NotoSansCJK", "app/assets/fonts/NotoSansJP-VariableFont_wght.ttf"))
        

    def _generate_filename(self, original_filename: str, extension: str) -> str:
        # stem_name = original_filename.rsplit(".", 1)[0]  
        # {Problem with latin-1 encoding, unable to handle edge-case languages. Thus, changed it to just document_summarized_timestamp.extnsion}
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        return f"document_summarized_{timestamp}.{extension}"
    

    def export_docx(self, original_filename: str, consolidated: ConsolidatedResult, model_results: list[ModelResult]) -> tuple[BytesIO, str]:
        document = Document()
        document.add_heading("Consolidated Summary", level = 0)
        document.add_paragraph(f"Document: {original_filename}")
        document.add_paragraph(clean_markdown(consolidated.summary))
        document.add_heading("Supporting Summaries", level = 1)
        
        for result in model_results:
            document.add_heading(f"Summary by: {result.model_name}", level = 2)
            document.add_paragraph(clean_markdown(result.summary))

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        return (buffer, self._generate_filename(original_filename, "docx"))


    def export_pdf(self, original_filename: str, consolidated: ConsolidatedResult, model_results: list[ModelResult]) -> tuple[BytesIO, str]:
        buffer = BytesIO()
        document = SimpleDocTemplate(buffer)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("TitleUnicode", parent = styles["Title"], fontName = "NotoSansCJK")
        heading1_style = ParagraphStyle("Heading1Unicode", parent = styles["Heading1"], fontName = "NotoSansCJK")
        heading2_style = ParagraphStyle("Heading2Unicode", parent = styles["Heading2"], fontName = "NotoSansCJK")
        normal_style = ParagraphStyle("NormalUnicode", parent = styles["Normal"], fontName = "NotoSansCJK", fontSize = 12, leading = 18)

        elements = []
        elements.append(Paragraph("<b>Consolidated Summary</b>", title_style))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(f"Document: {original_filename}", normal_style))
        elements.append(Spacer(1, 20))
        elements.append(Paragraph(clean_markdown(consolidated.summary), normal_style))
        elements.append(Spacer(1, 20))
        elements.append(Paragraph("<b>Supporting Summaries</b>", heading1_style))

        for result in model_results:
            elements.append(Paragraph(f"<b>Summary by: {result.model_name}</b>", heading2_style))
            elements.append(Paragraph(clean_markdown(result.summary), normal_style))
            elements.append(Spacer(1, 10))
        
        document.build(elements)
        buffer.seek(0)

        return (buffer, self._generate_filename(original_filename, "pdf"))